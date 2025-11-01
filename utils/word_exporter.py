from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import os

def export_to_word(title, subject, grid, word_list, font_name, filename, theme='modern', shape='square'):
    """
    Export puzzle to Word document with proper formatting for 6.5" x 9" page.
    
    Layout improvements:
    - Custom 6.5" x 9" page size
    - Title and subtitle centered on the puzzle grid
    - Word list on the left side with header aligned to puzzle grid
    - Increased word list column width for longer words
    - Centered content on page
    - White background throughout
    - Proper instruction width constraints
    """
    
    # Map non-standard fonts to Windows-compatible fonts
    font_mapping = {
        'Inter': 'Arial',
        'Inter (Modern)': 'Arial'
    }
    font_name = font_mapping.get(font_name, font_name)
    
    # Get theme colors
    theme_colors = get_theme_colors(theme)
    
    # Create new document
    doc = Document()
    
    # Set custom 6.5" x 9" page size and conditional margins
    for section in doc.sections:
        section.page_width = Inches(6.5)   # 6.5" width
        section.page_height = Inches(9.0)  # 9" height
        
        if shape == 'square':
            # Square shape: reduced margins (50% decrease on left/right)
            section.top_margin = Inches(1.5)   # 1.5" top margin
            section.bottom_margin = Inches(1.0)  # 1.0" bottom margin
            section.left_margin = Inches(0.25)   # 0.25" left margin (50% of 0.5")
            section.right_margin = Inches(0.375)  # 0.375" right margin (50% of 0.75")
        else:
            # Non-square shapes: equal margins for maximum wordlist space
            section.top_margin = Inches(0.3)   # Minimal margins for better space utilization
            section.bottom_margin = Inches(0.3)
            section.left_margin = Inches(0.3)  # Equal left margin
            section.right_margin = Inches(0.3)  # Equal right margin (same as left)
    
    # Set document background to white by setting paragraph styles
    doc.styles['Normal'].font.color.rgb = RGBColor.from_string('000000')  # Black text
    
    # Add title and subtitle to the header (for all shapes)
    if title or subject:
        add_title_and_subject_to_header(doc, title, subject, font_name, theme_colors)
    
    # Create main table with wordlist, spacer, and puzzle
    if shape == 'square':
        # Square shapes: 3 columns (wordlist, spacer, puzzle) for double spacing
        main_table = doc.add_table(rows=1, cols=3)
        main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        left_cell = main_table.cell(0, 0)    # Wordlist
        spacer_cell = main_table.cell(0, 1)  # Spacer
        right_cell = main_table.cell(0, 2)   # Puzzle
    else:
        # Non-square shapes: 3 columns (wordlist, spacer, puzzle) for proper spacing
        main_table = doc.add_table(rows=1, cols=3)
        main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        left_cell = main_table.cell(0, 0)    # Wordlist
        spacer_cell = main_table.cell(0, 1)  # Spacer
        right_cell = main_table.cell(0, 2)   # Puzzle
    
    if shape == 'square':
        # Square shape: with double spacing between columns
        left_cell.width = Inches(2.5)   # Wordlist width for square shapes
        spacer_cell.width = Inches(0.4)  # Double spacing for square shapes (2x larger)
        right_cell.width = Inches(2.35)  # Puzzle width for square shapes (reduced to accommodate spacer)
    else:
        # Non-square shapes: massive wordlist width with minimal spacing
        left_cell.width = Inches(5.1)   # Massive wordlist width for elderly users
        spacer_cell.width = Inches(0.15)  # Minimal spacer for visual separation
        right_cell.width = Inches(0.65)  # Adequate puzzle column space
    
    # Set table background to white and reduce padding
    for row in main_table.rows:
        for cell in row.cells:
            set_cell_background(cell, 'ffffff')
            # Reduce cell padding to maximize usable space
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            # Set minimal margins for cell content
            cell._tc.get_or_add_tcPr().append(OxmlElement('w:tcMar'))
            for margin_type in ['top', 'bottom', 'left', 'right']:
                margin = OxmlElement(f'w:{margin_type}')
                margin.set(qn('w:w'), '0')
                margin.set(qn('w:type'), 'dxa')
                cell._tc.get_or_add_tcPr().find(qn('w:tcMar')).append(margin)
    
    # Clear the spacer cell content
    spacer_cell.text = ''  # Empty spacer cell
    
    # Add puzzle to right cell (no titles - they're now on the page)
    add_puzzle_to_cell(right_cell, grid, font_name, theme_colors, shape)
    
    # Add word list to left cell (without header since it's now at the top)
    add_word_list_to_cell_without_header(left_cell, word_list, font_name, theme_colors, shape)
    
    # Add conditional spacing after main content
    spacer_paragraph2 = doc.add_paragraph()
    if shape == 'square':
        spacer_paragraph2.paragraph_format.space_after = Pt(3)  # Minimal spacing for square
    else:
        spacer_paragraph2.paragraph_format.space_after = Pt(2)  # Minimal spacing for non-square to fit one page
    
    # Conditional instructions placement based on shape
    if shape == 'square':
        # For square shapes: instructions go in footer (left aligned)
        add_instructions_to_footer(doc, font_name, theme_colors)
    else:
        # For non-square shapes: instructions go in footer (same as PDF format)
        add_instructions_to_footer_for_non_square(doc, font_name, theme_colors)
    
    # Save document
    normalized_path = os.path.abspath(filename)
    print(f"DEBUG word_exporter: Attempting to save to: {normalized_path}")
    print(f"DEBUG word_exporter: Directory exists: {os.path.exists(os.path.dirname(normalized_path))}")
    doc.save(normalized_path)

def set_cell_background(cell, color_hex):
    """Set the background color of a table cell to white."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def remove_table_borders(cell):
    """Remove borders from table cell to minimize gaps."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Remove top border
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'none')
    top_border.set(qn('w:sz'), '0')
    top_border.set(qn('w:space'), '0')
    top_border.set(qn('w:color'), 'auto')
    
    # Remove bottom border
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'none')
    bottom_border.set(qn('w:sz'), '0')
    bottom_border.set(qn('w:space'), '0')
    bottom_border.set(qn('w:color'), 'auto')
    
    # Remove left border
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'none')
    left_border.set(qn('w:sz'), '0')
    left_border.set(qn('w:space'), '0')
    left_border.set(qn('w:color'), 'auto')
    
    # Remove right border
    right_border = OxmlElement('w:right')
    right_border.set(qn('w:val'), 'none')
    right_border.set(qn('w:sz'), '0')
    right_border.set(qn('w:space'), '0')
    right_border.set(qn('w:color'), 'auto')
    
    # Create borders element and add all borders
    borders = OxmlElement('w:tcBorders')
    borders.append(top_border)
    borders.append(bottom_border)
    borders.append(left_border)
    borders.append(right_border)
    
    tcPr.append(borders)

def add_title_and_subject_to_header(doc, title, subject, font_name, theme_colors):
    """Add title and subject centered in the header (for all shapes)."""
    
    # Get the header from the first section
    header = doc.sections[0].header
    
    # Add title if provided
    if title:
        title_paragraph = header.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.font.name = font_name
        title_run.font.size = Pt(20)  # Match puzzle letters font size (20pt)
        title_run.bold = True  # Make title bold for emphasis
        title_run.font.color.rgb = RGBColor.from_string('000000')  # Black
        title_paragraph.paragraph_format.space_after = Pt(6)
    
    # Add subject if provided
    if subject:
        subject_paragraph = header.add_paragraph()
        subject_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subject_run = subject_paragraph.add_run(subject)
        subject_run.font.name = font_name
        subject_run.font.size = Pt(14)  # Subtitles = 14pt
        subject_run.font.color.rgb = RGBColor.from_string('000000')  # Black
        subject_paragraph.paragraph_format.space_after = Pt(8)

def add_puzzle_with_title_to_cell(cell, title, subject, grid, font_name, theme_colors, shape):
    """Add puzzle with title and subtitle centered on it to the specified cell."""
    
    # Add title if provided
    if title:
        title_paragraph = cell.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.font.name = font_name
        title_run.font.size = Pt(16)  # Titles = 16pt
        title_run.bold = True
        title_run.font.color.rgb = RGBColor.from_string('000000')  # Black
        title_paragraph.paragraph_format.space_after = Pt(4)  # Reduced spacing
    
    # Add subject if provided
    if subject:
        subject_paragraph = cell.add_paragraph()
        subject_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subject_run = subject_paragraph.add_run(subject)
        subject_run.font.name = font_name
        subject_run.font.size = Pt(14)  # Subtitles = 14pt
        subject_run.font.color.rgb = RGBColor.from_string('000000')  # Black
        subject_paragraph.paragraph_format.space_after = Pt(6)  # Reduced spacing
    
    # Add puzzle table
    add_puzzle_to_cell(cell, grid, font_name, theme_colors, shape)

def add_word_list_to_cell(cell, word_list, font_name, theme_colors):
    """Add word list to the specified cell with proper alignment."""
    
    # Add minimal padding to align wordlist title with puzzle grid top
    padding_paragraph = cell.add_paragraph()
    padding_paragraph.paragraph_format.space_after = Pt(2)  # Minimal padding for proper alignment
    
    # Add header
    header_paragraph = cell.add_paragraph()
    header_run = header_paragraph.add_run("Wordlist:")
    header_run.font.name = font_name
    header_run.font.size = Pt(16)  # Wordlist header = 16pt
    header_run.bold = True
    header_run.font.color.rgb = RGBColor.from_string('000000')  # Black
    header_paragraph.paragraph_format.space_after = Pt(3)  # Reduced spacing after wordlist title
    
    # Add words
    for word in word_list:
        # Capitalize and preserve spaces within words
        formatted_word = word.upper()
        
        word_paragraph = cell.add_paragraph()
        word_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        word_run = word_paragraph.add_run(formatted_word)
        word_run.font.name = font_name
        word_run.font.size = Pt(11)  # Words = 11pt
        word_run.font.color.rgb = RGBColor.from_string(theme_colors['text'])
        
        # Minimal spacing between words
        word_paragraph.paragraph_format.space_after = Pt(2)

def add_word_list_to_cell_without_header(cell, word_list, font_name, theme_colors, shape='square'):
    """Add word list to the specified cell with header aligned to puzzle grid top."""
    
    # Add "Wordlist:" header aligned with puzzle grid top
    header_paragraph = cell.add_paragraph()
    header_run = header_paragraph.add_run("Wordlist:")
    header_run.font.name = font_name
    if shape == 'square':
        header_run.font.size = Pt(16)  # Square shapes: 16pt (reduced from 20pt)
    else:
        header_run.font.size = Pt(16)  # Non-square shapes: 16pt (keep as is)
    header_run.bold = True  # Make wordlist header bold for emphasis
    header_run.font.color.rgb = RGBColor.from_string('000000')  # Black
    header_paragraph.paragraph_format.space_after = Pt(6)  # Reduced space between header and first word
    
    # Add words with optimized spacing to align bottom with puzzle grid
    for i, word in enumerate(word_list):
        # Capitalize and preserve spaces within words
        formatted_word = word.upper()
        
        word_paragraph = cell.add_paragraph()
        word_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        word_run = word_paragraph.add_run(formatted_word)
        word_run.font.name = font_name
        word_run.font.size = Pt(12)  # Keep 12pt for readability (elderly-friendly)
        word_run.font.color.rgb = RGBColor.from_string(theme_colors['text'])
        
        # Optimized spacing: minimal between words, no spacing after last word
        if i < len(word_list) - 1:
            word_paragraph.paragraph_format.space_after = Pt(1)  # Minimal spacing between words
        else:
            word_paragraph.paragraph_format.space_after = Pt(0)  # No spacing after last word

def add_puzzle_to_cell(cell, grid, font_name, theme_colors, shape='square'):
    """Add puzzle grid to the specified cell."""
    
    # Create puzzle table
    puzzle_table = cell.add_table(rows=len(grid), cols=len(grid[0]))
    puzzle_table.style = 'Table Grid'
    puzzle_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set conditional cell sizes based on shape - optimized for one-page fit
    if shape == 'square':
        # Square shapes: larger cells since we have more space
        cell_size = Inches(0.22)  # 0.22" cells for square shapes
    else:
        # Non-square shapes: smaller cells to ensure one-page fit
        cell_size = Inches(0.18)  # 0.18" cells for non-square shapes
    
    for row in puzzle_table.rows:
        row.height = cell_size
        for cell in row.cells:
            cell.width = cell_size
    
    # Fill table with grid data
    for i, row in enumerate(grid):
        for j, cell_content in enumerate(row):
            table_cell = puzzle_table.cell(i, j)
            table_cell.text = cell_content
            
            # Set cell background to white
            set_cell_background(table_cell, 'ffffff')
            
            # Center text in each cell
            for paragraph in table_cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = font_name
                    if shape == 'square':
                        run.font.size = Pt(20)  # Square shapes: 20pt puzzle letters (original)
                    else:
                        run.font.size = Pt(12)  # Non-square shapes: 12pt puzzle letters
                    run.bold = False  # Make puzzle letters not bold
                    run.font.color.rgb = RGBColor.from_string(theme_colors['text'])

def add_instructions_to_footer(doc, font_name, theme_colors):
    """Add instructions to the footer, formatted like the image."""
    
    # Get the footer
    footer = doc.sections[0].footer
    
    # Add instructions title
    title_paragraph = footer.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_paragraph.add_run("Instructions:")
    title_run.font.name = font_name
    title_run.font.size = Pt(11)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string('000000')  # Black
    
    # Add instruction lines (left-aligned, matching image format)
    instructions = [
        "Find all the words hidden in the puzzle above.",
        "Circle or highlight each word as you find it.",
        "Words can be read horizontally, vertically, or diagonally."
    ]
    
    for instruction in instructions:
        instruction_paragraph = footer.add_paragraph()
        instruction_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        instruction_paragraph.paragraph_format.space_after = Pt(0)  # No spacing between sentences
        instruction_run = instruction_paragraph.add_run(instruction)
        instruction_run.font.name = font_name
        instruction_run.font.size = Pt(11)
        instruction_run.font.color.rgb = RGBColor.from_string('000000')  # Black

def add_instructions_to_footer_for_non_square(doc, font_name, theme_colors):
    """Add instructions to the footer for non-square shapes, matching PDF format."""
    
    # Get the footer
    footer = doc.sections[0].footer
    
    # Add instructions title
    title_paragraph = footer.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_paragraph.add_run("Instructions:")
    title_run.font.name = font_name
    title_run.font.size = Pt(11)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string('000000')  # Black
    
    # Add instruction text with bold "Note:" (matching PDF format)
    instruction_paragraph = footer.add_paragraph()
    instruction_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    instruction_paragraph.paragraph_format.space_after = Pt(0)
    
    # Add main instruction text
    main_run = instruction_paragraph.add_run("Find all the words hidden in the puzzle above, circle or highlight each word as you find it. ")
    main_run.font.name = font_name
    main_run.font.size = Pt(11)
    main_run.font.color.rgb = RGBColor.from_string('000000')  # Black
    
    # Add "Note:" in bold
    note_bold_run = instruction_paragraph.add_run("Note: ")
    note_bold_run.font.name = font_name
    note_bold_run.font.size = Pt(11)
    note_bold_run.bold = True
    note_bold_run.font.color.rgb = RGBColor.from_string('000000')  # Black
    
    # Add rest of note text
    note_text_run = instruction_paragraph.add_run("Words may be read horizontally, vertically, or diagonally.")
    note_text_run.font.name = font_name
    note_text_run.font.size = Pt(11)
    note_text_run.font.color.rgb = RGBColor.from_string('000000')  # Black

def add_constrained_instructions(doc, font_name, theme_colors, max_width):
    """Add instructions constrained to a specific width, formatted like the image."""
    
    # Create a table to constrain the instructions width
    instructions_table = doc.add_table(rows=1, cols=1)
    instructions_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set the table width to match puzzle width
    instructions_cell = instructions_table.cell(0, 0)
    instructions_cell.width = max_width
    
    # Set cell background to white
    set_cell_background(instructions_cell, 'ffffff')
    
    # Add instructions title
    title_paragraph = instructions_cell.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_paragraph.add_run("Instructions:")
    title_run.font.name = font_name
    title_run.font.size = Pt(11)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string('000000')  # Black
    
    # Add single line instruction with bold "Note:"
    instruction_paragraph = instructions_cell.add_paragraph()
    instruction_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    instruction_paragraph.paragraph_format.space_after = Pt(0)
    
    # Add main instruction text
    main_run = instruction_paragraph.add_run("Find all the words hidden in the puzzle above, circle or highlight each word as you find it. ")
    main_run.font.name = font_name
    main_run.font.size = Pt(11)
    main_run.font.color.rgb = RGBColor.from_string('000000')  # Black
    
    # Add "Note: " in bold
    note_bold_run = instruction_paragraph.add_run("Note: ")
    note_bold_run.font.name = font_name
    note_bold_run.font.size = Pt(11)
    note_bold_run.bold = True
    note_bold_run.font.color.rgb = RGBColor.from_string('000000')  # Black
    
    # Add the rest of the note in regular font
    note_text_run = instruction_paragraph.add_run("Words may be read horizontally, vertically, or diagonally.")
    note_text_run.font.name = font_name
    note_text_run.font.size = Pt(11)
    note_text_run.bold = False
    note_text_run.font.color.rgb = RGBColor.from_string('000000')  # Black

def get_theme_colors(theme):
    """Get color scheme for the specified theme."""
    
    if theme == 'modern':
        return {
            'primary': '1e3a8a',        # Dark blue for titles (no #)
            'secondary': '3b82f6',       # Light blue for subtitles (no #)
            'text': '000000',            # Black text for puzzle and word list (no #)
            'background': 'ffffff',      # White background (no #)
            'grid': 'd1d5db'             # Light grey grid lines (no #)
        }
    elif theme == 'cozy':
        return {
            'primary': '166534',         # Dark green (no #)
            'secondary': '22c55e',       # Light green (no #)
            'text': '000000',            # Black text (no #)
            'background': 'ffffff',      # White background (no #)
            'grid': 'd1d5db'             # Light grey (no #)
        }
    elif theme == 'playful':
        return {
            'primary': '7c3aed',         # Purple (no #)
            'secondary': 'f97316',       # Orange (no #)
            'text': '000000',            # Black text (no #)
            'background': 'ffffff',      # White background (no #)
            'grid': 'd1d5db'             # Light grey (no #)
        }
    else:
        # Default to modern with black text on white background
        return {
            'primary': '1e3a8a',        # Dark blue for titles (no #)
            'secondary': '3b82f6',       # Light blue for subtitles (no #)
            'text': '000000',            # Black text for puzzle and word list (no #)
            'background': 'ffffff',      # White background (no #)
            'grid': 'd1d5db'             # Light grey grid lines (no #)
        }